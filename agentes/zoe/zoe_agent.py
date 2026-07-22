"""
ZOE - Secretaria IA de Ruth Inmobiliaria
Gestiona documentos, contratos, facturas, memoria, y conversa con IA real (OpenRouter).
"""
import os
import json
import re
import io
import base64
from datetime import datetime

import fitz  # PyMuPDF
import requests
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE_DIR = os.environ.get("DATA_DIR", os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))))
CONFIG_FILE = os.path.join(BASE_DIR, "config", "config.json")


def load_config():
    if os.path.exists(CONFIG_FILE):
        with open(CONFIG_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    return {}


CONFIG = load_config()
OPENROUTER_API_KEY = CONFIG.get("OPENROUTER_API_KEY", "")
SUPABASE_URL = CONFIG.get("SUPABASE_URL", "")
SUPABASE_KEY = CONFIG.get("SUPABASE_KEY", "")


class ZOEAgent:
    """Agente ZOE: Secretaria IA de Ruth Inmobiliaria con OpenRouter."""

    def __init__(self):
        self.config = load_config()
        self.api_key = self.config.get("OPENROUTER_API_KEY", "")
        self.base_url = self.config.get("OPENROUTER_BASE_URL", "https://openrouter.ai/api/v1")
        self.model = self.config.get("MODEL", "deepseek/deepseek-chat")
        
        # System prompt de ZOE
        self.system_prompt = """Eres ZOE, la secretaria IA de Ruth Inmobiliaria (RB Inmobiliaria).
Trabajas para Ruth Blanco, agente inmobiliario en Castelldefels, Gavà y Viladecans.

TUS CAPACIDADES:
- Redactar contratos de alquiler y arras
- Generar facturas profesionales
- Leer y extraer información de PDFs y documentos
- Guardar y recuperar memoria (información de clientes, propiedades, etc.)
- Generar informes de mercado, actividad y resúmenes
- Responder preguntas sobre la empresa y sus procesos
- Buscar información en internet cuando sea necesario

INSTRUCCIONES:
- Eres PROFESIONAL pero CERCANA, como una secretaria de confianza
- Si te piden CREAR un contrato/factura, usa las funciones crear_contrato() o generar_factura()
- Si te piden LEER un documento, usa leer_documento()
- Si te piden GUARDAR o RECORDAR algo, usa guardar_memoria()
- Si preguntan algo general o no cubierto arriba, responde directamente con tu conocimiento
- Respuestas en español, claras y directas
- Si no sabes algo, dilo honestamente"""

    def _call_llm(self, messages, temperature=0.3, max_tokens=2048):
        """Llama a OpenRouter con deepseek-v4-flash."""
        if not self.api_key or "AQUI" in self.api_key:
            return "⚠️ ZOE no tiene conexión con la IA. Configura OPENROUTER_API_KEY en config.json."
        try:
            headers = {
                "Authorization": f"Bearer {self.api_key}",
                "Content-Type": "application/json",
                "HTTP-Referer": "https://ruthinmobiliaria.com",
            }
            payload = {
                "model": self.model,
                "messages": [{"role": "system", "content": self.system_prompt}] + messages,
                "temperature": temperature,
                "max_tokens": max_tokens,
            }
            resp = requests.post(f"{self.base_url}/chat/completions", json=payload, headers=headers, timeout=30)
            if resp.status_code == 200:
                return resp.json()["choices"][0]["message"]["content"]
            return f"⚠️ Error LLM: HTTP {resp.status_code}"
        except Exception as e:
            return f"⚠️ Error llamando a la IA: {e}"

    def _log_conversacion(self, agente: str, mensaje: str, respuesta: str = None):
        """Guarda una conversación en Supabase."""
        try:
            url = f"{SUPABASE_URL}/rest/v1/conversaciones"
            headers = {
                "apikey": SUPABASE_KEY,
                "Authorization": f"Bearer {SUPABASE_KEY}",
                "Content-Type": "application/json",
            }
            payload = {
                "agente": agente,
                "usuario": "Ruth",
                "mensaje": mensaje[:1000],
                "respuesta": respuesta[:2000] if respuesta else "",
            }
            requests.post(url, headers=headers, json=payload, timeout=5)
        except:
            pass

    def _log_accion(self, agente: str, accion: str, resultado: str = "ok", error: str = None):
        """Guarda un log de acción en Supabase."""
        try:
            import requests as _req
            url = f"{SUPABASE_URL}/rest/v1/logs_agentes"
            headers = {
                "apikey": SUPABASE_KEY,
                "Authorization": f"Bearer {SUPABASE_KEY}",
                "Content-Type": "application/json",
            }
            payload = {
                "agente": agente,
                "accion": accion,
                "resultado": resultado[:500],
                "error": error[:500] if error else None,
            }
            _req.post(url, headers=headers, json=payload, timeout=5)
        except:
            pass

    def leer_pdf(self, ruta: str) -> str:
        """Extrae texto de un PDF."""
        try:
            doc = fitz.open(ruta)
            texto = ""
            for pagina in doc:
                texto += pagina.get_text() + "\n"
            doc.close()
            return texto.strip()
        except Exception as e:
            return f"Error leyendo PDF: {e}"

    def leer_documento(self, ruta: str) -> str:
        """Lee cualquier documento soportado."""
        ext = os.path.splitext(ruta)[1].lower()
        if ext == ".pdf":
            return self.leer_pdf(ruta)
        elif ext in (".txt", ".md", ".csv"):
            with open(ruta, "r", encoding="utf-8") as f:
                return f.read()
        elif ext in (".docx", ".doc"):
            doc = Document(ruta)
            return "\n".join([p.text for p in doc.paragraphs])
        else:
            return f"Formato no soportado: {ext}"

    def crear_contrato(self, datos: dict) -> str:
        """Genera un contrato DOCX a partir de datos."""
        plantilla_dir = os.path.join(BASE_DIR, "CONTRATOS")
        os.makedirs(plantilla_dir, exist_ok=True)

        doc = Document()

        # Título
        titulo = doc.add_heading("CONTRATO DE ARRENDAMIENTO", level=0)
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

        fecha = datetime.now().strftime("%d de %B de %Y")
        doc.add_paragraph(
            f"Fecha: {fecha}",
            style="Normal"
        )

        # Arrendador
        if "arrendador" in datos:
            p = doc.add_paragraph()
            p.add_run("ARRENDADOR: ").bold = True
            p.add_run(datos.get("arrendador", ""))

        # Arrendatario
        if "arrendatario" in datos:
            p = doc.add_paragraph()
            p.add_run("ARRENDATARIO: ").bold = True
            p.add_run(datos.get("arrendatario", ""))

        # Propiedad
        if "propiedad" in datos:
            p = doc.add_paragraph()
            p.add_run("PROPIEDAD: ").bold = True
            p.add_run(datos.get("propiedad", ""))

        # DNI
        if "dni" in datos:
            p = doc.add_paragraph()
            p.add_run("DNI/CIF: ").bold = True
            p.add_run(datos.get("dni", ""))

        # Cuerpo del contrato
        doc.add_heading("CLÁUSULAS", level=1)
        clausulas = datos.get("clausulas", [
            "El arrendador cede en alquiler el inmueble descrito.",
            "El arrendatario se compromete al pago puntual del alquiler.",
            "La duración del contrato será de acuerdo a la legislación vigente.",
            "Los gastos de comunidad serán a cargo del arrendatario.",
            "Cualquier modificación requerirá acuerdo por escrito de ambas partes."
        ])

        for i, clausula in enumerate(clausulas, 1):
            doc.add_paragraph(f"  {i}. {clausula}")

        # Importe
        if "importe" in datos:
            p = doc.add_paragraph()
            p.add_run("IMPORTE MENSUAL: ").bold = True
            run = p.add_run(datos.get("importe", "€"))
            run.font.color.rgb = RGBColor(0x00, 0x70, 0xC0)

        # Firma
        doc.add_paragraph("\n\n\n")
        tabla = doc.add_table(rows=2, cols=2)
        tabla.cell(0, 0).text = "Firma Arrendador"
        tabla.cell(0, 1).text = "Firma Arrendatario"
        tabla.cell(1, 0).text = "________________________"
        tabla.cell(1, 1).text = "________________________"

        nombre = datos.get("arrendatario", "contrato").replace(" ", "_")
        ruta_out = os.path.join(plantilla_dir, f"CONTRATO_{nombre}_{datetime.now().strftime('%Y%m%d')}.docx")
        doc.save(ruta_out)
        return f"Contrato generado: {ruta_out}"

    def generar_factura(self, datos: dict) -> str:
        """Genera una factura DOCX."""
        factura_dir = os.path.join(BASE_DIR, "FACTURAS_IA_PROCESAR")
        os.makedirs(factura_dir, exist_ok=True)

        doc = Document()

        # Encabezado
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run("FACTURA\n")
        run.bold = True
        run.font.size = Pt(16)

        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p2.add_run(f"Fecha: {datetime.now().strftime('%d/%m/%Y')}\n")
        if "numero" in datos:
            p2.add_run(f"Nº: {datos['numero']}")

        # Destinatario
        doc.add_heading("DATOS DEL CLIENTE", level=1)
        campos = ["nombre", "direccion", "nif", "email", "telefono"]
        for campo in campos:
            if campo in datos:
                p = doc.add_paragraph()
                p.add_run(f"{campo.upper()}: ").bold = True
                p.add_run(str(datos[campo]))

        # Líneas
        doc.add_heading("CONCEPTOS", level=1)
        tabla = doc.add_table(rows=1, cols=4)
        hdr = tabla.rows[0].cells
        for i, h in enumerate(["Concepto", "Cantidad", "Precio", "Total"]):
            hdr[i].text = h
            hdr[i].paragraphs[0].runs[0].bold = True

        lineas = datos.get("lineas", [])
        total = 0
        for linea in lineas:
            row_cells = tabla.add_row().cells
            row_cells[0].text = linea.get("concepto", "")
            row_cells[1].text = str(linea.get("cantidad", 1))
            precio = linea.get("precio", 0)
            row_cells[2].text = f"{precio:.2f} €"
            subtotal = linea.get("cantidad", 1) * precio
            row_cells[3].text = f"{subtotal:.2f} €"
            total += subtotal

        # Total
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.add_run("TOTAL: ").bold = True
        run = p.add_run(f"{total:.2f} €")
        run.font.size = Pt(14)

        nombre = datos.get("nombre", "factura").replace(" ", "_")
        ruta_out = os.path.join(factura_dir, f"FACTURA_{nombre}_{datetime.now().strftime('%Y%m%d')}.docx")
        doc.save(ruta_out)
        return f"Factura generada: {ruta_out}"

    def guardar_memoria(self, key: str, value: str) -> str:
        """Guarda un dato en Supabase (memoria_agentes)."""
        import requests as _req
        url = f"{SUPABASE_URL}/rest/v1/memoria_agentes"
        headers = {
            "apikey": SUPABASE_KEY,
            "Authorization": f"Bearer {SUPABASE_KEY}",
            "Content-Type": "application/json",
            "Prefer": "resolution=merge-duplicates",
        }
        payload = {
            "agente": "ZOE",
            "clave": key,
            "valor": value,
            "updated_at": datetime.now().isoformat(),
        }
        try:
            r = _req.post(url, headers=headers, json=payload, timeout=10)
            if r.status_code in (200, 201):
                return f"🧠 ZOE ha guardado: {key}"
        except:
            pass
        # Fallback a disco
        mem_dir = os.path.join(BASE_DIR, "output", "memoria")
        os.makedirs(mem_dir, exist_ok=True)
        with open(os.path.join(mem_dir, f"{key.replace(' ', '_')}.json"), "w", encoding="utf-8") as f:
            json.dump({"clave": key, "valor": value}, f)
        return f"🧠 ZOE ha guardado: {key} (disco)"

    def leer_memoria(self, key: str = None) -> str:
        """Lee de Supabase memoria_agentes."""
        import requests as _req
        headers = {
            "apikey": SUPABASE_KEY,
            "Authorization": f"Bearer {SUPABASE_KEY}",
        }
        try:
            if key:
                url = f"{SUPABASE_URL}/rest/v1/memoria_agentes?agente=eq.ZOE&clave=eq.{key}&select=clave,valor"
            else:
                url = f"{SUPABASE_URL}/rest/v1/memoria_agentes?agente=eq.ZOE&select=clave,valor&order=updated_at.desc&limit=20"
            r = _req.get(url, headers=headers, timeout=10)
            if r.status_code == 200 and r.json():
                datos = r.json()
                if key:
                    d = datos[0]
                    return f"📝 {d['clave']}: {d['valor']}"
                res = "🧠 Memoria de ZOE:\n"
                for d in datos:
                    v = d['valor']
                    if isinstance(v, str) and len(v) > 60:
                        v = v[:60] + "..."
                    res += f"  • {d['clave']}: {v}\n"
                return res
        except:
            pass
        return "No encontré nada en memoria."

    def generar_informe(self, tipo: str, datos: dict = None) -> str:
        """Genera un informe en formato texto."""
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M")
        informe = f"=== INFORME: {tipo.upper()} ===\nFecha: {timestamp}\n"

        if tipo == "mercado":
            informe += f"\nZonas: {', '.join(datos.get('zonas', ['Castelldefels', 'Gavà']))}\n"
            informe += f"Captaciones activas: {datos.get('captaciones', 0)}\n"
            informe += f"Consultas recibidas: {datos.get('consultas', 0)}\n"
        elif tipo == "actividad":
            informe += f"\nTareas completadas: {datos.get('completadas', 0)}\n"
            informe += f"En espera: {datos.get('pendientes', 0)}\n"
        elif tipo == "resumen":
            informe += "\nResumen del día generado por ZOE.\n"

        if datos:
            informe += f"\nDatos adicionales:\n{json.dumps(datos, indent=2, default=str, ensure_ascii=False)}"

        return informe

    def execute(self, task: str) -> str:
        """Punto de entrada principal. Usa IA para entender y ejecutar tareas."""
        task_lower = task.lower()

        # Tareas detectables por palabras clave → funciones concretas
        if "contrato" in task_lower or "arras" in task_lower:
            return self.crear_contrato({"arrendador": "Ruth Blanco", "arrendatario": task})

        elif "factura" in task_lower:
            return self.generar_factura({"nombre": task})

        elif "leer" in task_lower or "pdf" in task_lower or "documento" in task_lower:
            # Extraer posible ruta
            import re as _re
            rutas = _re.findall(r'[A-Z]:(?:\\\\[\\w\\s.()-]+)+', task)
            if rutas:
                return self.leer_documento(rutas[0])
            return "Indica la ruta del archivo que quieres leer (ej: D:\\...\\documento.pdf)"

        elif "guardar" in task_lower or "recuerda" in task_lower:
            partes = task.split(":", 1)
            if len(partes) > 1:
                return self.guardar_memoria("nota", partes[1].strip())
            return self.guardar_memoria("nota", task.replace("guarda", "").replace("recuerda", "").strip())

        elif "informe" in task_lower:
            return self.generar_informe("resumen")

        elif "memoria" in task_lower or "buscar" in task_lower:
            return self.leer_memoria()

        # Todo lo demás → IA real
        else:
            respuesta = self._call_llm([
                {"role": "user", "content": task}
            ])
            self._log_conversacion("ZOE", task, respuesta)
            return respuesta